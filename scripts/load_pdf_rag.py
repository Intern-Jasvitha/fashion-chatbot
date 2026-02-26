#!/usr/bin/env python3
"""
Universal Knowledge Base Ingestion
Handles FAQ + Guides + Training Manuals + Decision Logic
"""

import argparse
import logging
import os
import sys
import uuid
import re
import time
from datetime import datetime, timezone
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(PROJECT_ROOT)

import httpx
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams, OptimizersConfigDiff

from app.core.config import get_settings
from app.core.qdrant import create_qdrant_client

DATA_DIR = PROJECT_ROOT / "data" / "Pdfs1"
BATCH_EMBED_SIZE = 32

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# -----------------------------------------------------------------------------
# EMBEDDINGS
# -----------------------------------------------------------------------------

class EmbeddingClient:
    def __init__(self, url: str):
        self.url = url.rstrip("/")

    def embed_batch(self, texts):
        with httpx.Client(timeout=120.0) as client:
            r = client.post(self.url, json={"text": texts})
            r.raise_for_status()
            data = r.json()

        if "embedding" in data:
            emb = data["embedding"]
            return emb if isinstance(emb[0], list) else [emb]
        if "embeddings" in data:
            return data["embeddings"]
        if "data" in data:
            return [d["embedding"] for d in data["data"]]

        raise ValueError("Unknown embedding response")

# -----------------------------------------------------------------------------
# ROBUST DOCX LOADER
# -----------------------------------------------------------------------------

def load_docx_structured(path: Path):
    doc = DocxDocument(path)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts)
    return [Document(page_content=full_text, metadata={"source": str(path)})]

def load_document(path: Path):
    if path.suffix.lower() == ".pdf":
        return PyPDFLoader(str(path)).load()
    if path.suffix.lower() == ".docx":
        return load_docx_structured(path)
    raise ValueError("Unsupported file")

# -----------------------------------------------------------------------------
# SPLITTERS
# -----------------------------------------------------------------------------

FAQ_PATTERN = re.compile(
    r"(?:Q[:\-]\s*)(?P<q>.*?)(?:\n|\r\n)+(?:A[:\-]\s*)(?P<a>.*?)(?=\nQ[:\-]|\Z)",
    re.DOTALL | re.IGNORECASE,
)

HEADING_SPLIT = re.compile(r"\n(?=[A-Z][^\n]{3,80}\n)", re.MULTILINE)

KNOWLEDGE_BLOCK_SPLIT = re.compile(
    r"\n(?=("
    r"\d+\.\s|"               
    r"rule set\s*\d*|"
    r"if\s+.*?\nthen|"
    r"internal checklist|"
    r"tone:|tone rules|"
    r"chatbot rule|chatbot must|chatbot behavior"
    r"))",
    re.IGNORECASE
)

# -----------------------------------------------------------------------------
# CHUNKING LOGIC
# -----------------------------------------------------------------------------

def semantic_chunk_documents(docs, source, filename):
    now = datetime.now(timezone.utc).isoformat()
    chunks = []

    for doc in docs:
        text = doc.page_content.strip()

        # ---------- FAQ ----------
        for i, m in enumerate(FAQ_PATTERN.finditer(text)):
            chunks.append({
                "id": str(uuid.uuid4()),
                "content": m.group("a").strip(),
                "metadata": {
                    "title": m.group("q").strip(),
                    "type": "faq",
                    "source": source,
                    "filename": filename,
                    "chunk_index": i,
                    "indexed_at": now,
                },
            })

        # ---------- KNOWLEDGE MANUAL ----------
        knowledge_blocks = KNOWLEDGE_BLOCK_SPLIT.split(text)
        if len(knowledge_blocks) > 5:
            for i, block in enumerate(knowledge_blocks):
                block = block.strip()
                if len(block) < 80:
                    continue

                title = block.split("\n")[0][:120]

                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": block,
                    "metadata": {
                        "title": title,
                        "type": "knowledge_rule",
                        "source": source,
                        "filename": filename,
                        "chunk_index": i,
                        "indexed_at": now,
                    },
                })
            continue

        # ---------- HEADING SECTIONS ----------
        sections = HEADING_SPLIT.split(text)
        meaningful = 0

        for i, section in enumerate(sections):
            section = section.strip()
            if len(section) < 120:
                continue

            meaningful += 1
            title = section.split("\n")[0][:120]

            chunks.append({
                "id": str(uuid.uuid4()),
                "content": section,
                "metadata": {
                    "title": title,
                    "type": "section",
                    "source": source,
                    "filename": filename,
                    "chunk_index": i,
                    "indexed_at": now,
                },
            })

        # ---------- FALLBACK ----------
        if meaningful == 0:
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if len(p.strip()) > 150]
            for i, para in enumerate(paragraphs):
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": para,
                    "metadata": {
                        "title": para[:80],
                        "type": "paragraph",
                        "source": source,
                        "filename": filename,
                        "chunk_index": i,
                        "indexed_at": now,
                    },
                })

    return chunks

# -----------------------------------------------------------------------------
# COLLECTION
# -----------------------------------------------------------------------------

def ensure_collection(client: QdrantClient, name: str, dim: int, recreate: bool):
    collections = [c.name for c in client.get_collections().collections]

    if recreate and name in collections:
        client.delete_collection(name)

    if name not in collections or recreate:
        client.create_collection(
            collection_name=name,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
            optimizers_config=OptimizersConfigDiff(indexing_threshold=0),
        )
        logger.info("Created collection")

def wait_for_indexing(client: QdrantClient, collection: str):
    logger.info("Building vector index...")

    client.update_collection(collection_name=collection, optimizers_config=OptimizersConfigDiff(indexing_threshold=1))

    for _ in range(60):
        info = client.get_collection(collection)
        logger.info("Index progress %d/%d", info.indexed_vectors_count, info.points_count)
        if info.indexed_vectors_count >= info.points_count and info.points_count > 0:
            logger.info("Semantic search READY")
            return
        time.sleep(1)

# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--recreate", action="store_true")
    args = parser.parse_args()

    settings = get_settings()
    client = create_qdrant_client()
    embedder = EmbeddingClient(settings.EMBEDDING_URL)

    ensure_collection(client, settings.QDRANT_COLLECTION_NAME, settings.EMBEDDING_DIMENSION, args.recreate)

    files = list(DATA_DIR.glob("*.pdf")) + list(DATA_DIR.glob("*.docx"))
    total = 0

    for file in files:
        logger.info("Processing %s", file.name)

        docs = load_document(file)
        chunks = semantic_chunk_documents(docs, str(file), file.name)

        if not chunks:
            logger.warning("No content extracted from %s", file.name)
            continue

        texts = [c["content"] for c in chunks]
        vectors = []

        for i in range(0, len(texts), BATCH_EMBED_SIZE):
            vectors.extend(embedder.embed_batch(texts[i:i+BATCH_EMBED_SIZE]))

        points = [PointStruct(id=c["id"], vector=v, payload={"content": c["content"], **c["metadata"]}) for c, v in zip(chunks, vectors)]
        client.upsert(settings.QDRANT_COLLECTION_NAME, points)

        total += len(points)
        logger.info("Inserted %d chunks", len(points))

    wait_for_indexing(client, settings.QDRANT_COLLECTION_NAME)
    logger.info("DONE — %d knowledge chunks indexed", total)

if __name__ == "__main__":
    main()
